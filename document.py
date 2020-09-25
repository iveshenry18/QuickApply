import docx
import os

class Document():
    """Represents a generic document"""
    def __init__(self, job, pmt='>>> '):
        self.pmt = pmt
        self.job = job
        self.doc = self.init_doc()


    def find_and_replace(self, tag, text):
        """Finds tag in self.doc and replaces with text
        
        Taken from https://stackoverflow.com/questions/34779724/python-docx-replace-string-in-paragraph-while-keeping-style
        """
        for p in self.doc.paragraphs:
            if tag in p.text:
                text = p.text.replace(tag, text)
                style = p.style
                p.text = text
                p.style = style

    def init_doc(self):
        return docx.Document()
